import sys, os
from docx import Document
from docx.shared import Inches
from pdf2image import convert_from_path
from tkinter import Tk, filedialog, messagebox

def main():
    if getattr(sys, '_MEIPASS', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(__file__)

    poppler_path = os.path.join(base_path, "poppler", "Library", "bin")

    # Hide tkinter root window
    Tk().withdraw()

    # Ask user for folder
    folder_path = filedialog.askdirectory(title="Select Folder Containing PDFs and JPGs")
    if not folder_path:
        messagebox.showinfo("Cancelled", "No folder selected. Exiting.")
        return

    doc = Document()
    processed_files = 0

    for filename in sorted(os.listdir(folder_path)):
        lower = filename.lower()
        file_path = os.path.join(folder_path, filename)

        # --- Handle PDF files ---
        if lower.endswith(".pdf"):
            print(f"Processing PDF: {filename}")
            doc.add_heading(filename, level=2)
            try:
                # Use the line below to use Poppler bundled in the same folder as script for .exe
                #images = convert_from_path(file_path, dpi=200, poppler_path=poppler_path)
                # Use this line to use Poppler installed system-wide C:\poppler-25.07.0\Library\bin\pdftoppm.exe
                images = convert_from_path(file_path, dpi=200)
                for page_num, img in enumerate(images, start=1):
                    img_path = os.path.join(folder_path, f"temp_page_{page_num}.jpg")
                    img.save(img_path, "JPEG")
                    doc.add_picture(img_path, width=Inches(5.5), height=Inches(7))
                    os.remove(img_path)
            except Exception as e:
                doc.add_paragraph(f"[Error processing {filename}: {e}]")
            doc.add_page_break()
            processed_files += 1

        # --- Handle JPG/JPEG files ---
        elif lower.endswith(".jpg") or lower.endswith(".jpeg"):
            print(f"Adding Image: {filename}")
            doc.add_heading(filename, level=2)
            try:
                doc.add_picture(file_path, width=Inches(5.5), height=Inches(7))
                doc.add_page_break()
                processed_files += 1
            except Exception as e:
                doc.add_paragraph(f"[Error adding {filename}: {e}]")

    # --- Save output ---
    if processed_files == 0:
        messagebox.showwarning("No Files Found", "No PDFs or JPGs found in the selected folder.")
        return

    output_folder = os.path.join(folder_path, "Output")
    os.makedirs(output_folder, exist_ok=True)

    output_path = os.path.join(output_folder, "Combined_PDFs_and_Images.docx")
    doc.save(output_path)

    messagebox.showinfo("Success ✅", f"Word document created:\n{output_path}")

if __name__ == "__main__":
    main()
